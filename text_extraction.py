import tempfile
import re
import requests
import pypdfium2 as pdfium
from docx import Document
import nltk
nltk.download('punkt')
nltk.download('punkt_tab')
nltk.download('stopwords')
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import unicodedata

def get_text_from_pdf(pdf_path: str) -> str:
    """
    Extracts text from a PDF file using pypdfium2
    """
    try:
        pdf = pdfium.PdfDocument(pdf_path)
        full_text = []
        for page_num in range(len(pdf)):
            try:
                text = pdf[page_num].get_textpage().get_text_range().strip()
                if text:
                    full_text.append(text)
            except Exception as e:
                print(f"Error extracting text from page {page_num}: {e}")
        return "\n".join(full_text)
    except Exception as e:
        return f"Error extracting text from PDF: {e}"
    
def get_text_from_docx(docx_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.
    """
    try:
        doc = Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            try:
                text = para.text.strip()
                if text:
                    full_text.append(text)
            except Exception as e:
                print(f"Error extracting text from paragraph: {e}")
        return "\n".join(full_text)
    except Exception as e:
        return f"Error extracting text from DOCX: {e}"
    
def get_text_from_txt(txt_path: str) -> str:
    """
    Reads and returns text from a plain TXT file.
    """
    encodings = ['utf-8', 'utf-16', 'ISO-8859-1', 'cp1252']
    for encoding in encodings:
        try:
            with open(txt_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError as e:
            continue
    return f"Error extracting text from TXT: Unable to decode file using available encodings."
    
def remove_non_valids(text: str) -> str:
    """
    Removes unwanted characters, keeping letters, digits, and math symbols.
    """
    math_symbols = set("=+-*/^()[]{}<>≤≥≠∑∫√∞≈→←↔∂∆θλπμσϕϵαβγδ")
    return "".join(
        char for char in text if unicodedata.category(char).startswith("L") or 
        char in math_symbols or 
        char.isalnum()
    )


def preprocess_text(text: str) -> str:
    """
    Normalizes, tokenizes, filters, and cleans input text.
    """
    text = re.sub(r'[\x00-\x1F\x7F]', ' ', text)
    text = unicodedata.normalize("NFKD", text)
    text = text.lower()
    tokens = word_tokenize(text)
    cleaned_tokens = [remove_non_valids(token) for token in tokens]
    stop_words = set(stopwords.words("english"))
    filtered_tokens = [token for token in cleaned_tokens if token and token not in stop_words]
    try:
        filtered_tokens = filtered_tokens[:len(tokens) - 1 - tokens[::-1].index("references")]
    except ValueError:
        pass
    return " ".join(filtered_tokens)